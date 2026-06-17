import re
import json
import pickle
import numpy as np
from pathlib import Path
from backend.config import settings


class RAGService:
    def __init__(self):
        self._st_model = None
        self.embedding_model = None
        self.chunks: list[str] = []
        self._embeddings: np.ndarray | None = None
        self._init_embeddings()
        self._save_dir = settings.DATA_DIR / "rag_index"
        self._save_dir.mkdir(parents=True, exist_ok=True)
        self._load_index()

    def _init_embeddings(self):
        import os
        os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
        os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
        try:
            from sentence_transformers import SentenceTransformer
            model_path = None
            for p in [
                Path(os.path.expanduser("~")) / ".cache/huggingface/hub/models--BAAI--bge-base-zh-v1.5/snapshots",
                Path(os.path.expanduser("~")) / ".cache/huggingface/hub/models--BAAI--bge-small-zh-v1.5/snapshots",
            ]:
                if p.exists():
                    snaps = [d for d in p.iterdir() if d.is_dir()]
                    for s in snaps:
                        if (s / "modules.json").exists():
                            model_path = str(s)
                            break
                    if model_path:
                        break
            if not model_path:
                model_path = settings.EMBEDDING_MODEL
            self._st_model = SentenceTransformer(model_path, device="cpu")
            self._embed = lambda texts: self._st_model.encode(
                texts, normalize_embeddings=True, show_progress_bar=False
            )
        except Exception:
            self._embed = None

    def _save_index(self):
        base = self._save_dir
        base.mkdir(parents=True, exist_ok=True)
        (base / "chunks.json").write_text(json.dumps(self.chunks, ensure_ascii=False), encoding="utf-8")
        if self._embeddings is not None:
            np.save(base / "embeddings.npy", self._embeddings)

    def _load_index(self):
        base = self._save_dir
        f = base / "chunks.json"
        if f.exists():
            self.chunks = json.loads(f.read_text(encoding="utf-8"))
        emb_f = base / "embeddings.npy"
        if emb_f.exists():
            self._embeddings = np.load(emb_f)

    async def add_document(self, file_path: str | Path) -> dict:
        file_path = Path(file_path)
        chunks = await self._chunk_document(file_path)
        if not chunks:
            return {"name": file_path.name, "chunks": 0}

        import asyncio

        old_len = len(self.chunks)
        self.chunks.extend(chunks)

        if self._embed is not None:
            new_embs = await asyncio.to_thread(self._embed, chunks)
            if self._embeddings is not None:
                self._embeddings = np.vstack([self._embeddings, new_embs])
            else:
                self._embeddings = new_embs

        self._save_index()
        return {"name": file_path.name, "chunks": len(chunks)}

    async def add_qa_pair(self, question: str, answer: str):
        text = f"问：{question}\n答：{answer}"
        import asyncio
        self.chunks.append(text)
        if self._embed is not None:
            new_embs = await asyncio.to_thread(self._embed, [text])
            if self._embeddings is not None:
                self._embeddings = np.vstack([self._embeddings, new_embs])
            else:
                self._embeddings = new_embs
        self._save_index()


    async def _chunk_document(self, file_path: Path) -> list[str]:
        text = ""
        suffix = file_path.suffix.lower()

        try:
            if suffix in (".txt", ".md"):
                text = file_path.read_text("utf-8")

            elif suffix == ".pdf":
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(file_path))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                except ImportError:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(str(file_path))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)

            elif suffix == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))
                    parts.append("\n".join(rows))
                text = "\n".join(parts)

            elif suffix == ".xlsx":
                import pandas as pd
                df = pd.read_excel(str(file_path))
                if "attraction_name" in df.columns and "attraction_content" in df.columns:
                    texts = []
                    for name, group in df.groupby("attraction_name"):
                        row = group.iloc[0]
                        content = row.get("attraction_content", "")
                        if pd.isna(content):
                            continue
                        content = re.sub(r"<[^>]+>", "", str(content))
                        atype = row.get("attraction_type", "")
                        line = f"【{name}】"
                        if atype and not pd.isna(atype):
                            line += f"\n类型：{atype}"
                        line += f"\n{content.strip()}"
                        texts.append(line)
                    text = "\n\n".join(texts)
                else:
                    text = str(df.to_string(index=False))

        except Exception:
            pass

        if not text.strip():
            return [f"知识块: {file_path.name} 的内容"]

        chunks = []
        paragraphs = re.split(r'\n\s*\n', text)
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 1 <= settings.CHUNK_SIZE:
                current = (current + "\n" + para) if current else para
            else:
                if current:
                    chunks.append(current.strip())
                if len(para) > settings.CHUNK_SIZE:
                    for i in range(0, len(para), settings.CHUNK_SIZE - settings.CHUNK_OVERLAP):
                        chunk = para[i:i + settings.CHUNK_SIZE].strip()
                        if chunk:
                            chunks.append(chunk)
                else:
                    current = para

        if current:
            chunks.append(current.strip())

        return chunks if chunks else [text.strip()[:settings.CHUNK_SIZE]]

    async def retrieve(self, query: str, top_k: int = None) -> list[str]:
        """检索相关知识，并使用 LLM 优化检索结果"""
        if top_k is None:
            top_k = settings.TOP_K

        # BGE neural search
        if self._embed is not None and self._embeddings is not None and len(self.chunks) > 0:
            import asyncio
            qvec = await asyncio.to_thread(self._embed, [query])
            qvec = qvec[0]
            sims = (self._embeddings @ qvec).flatten()
            idxs = np.argsort(sims)[::-1][:top_k]
            results = [self.chunks[i] for i in idxs if sims[i] > 0.3]
            if results:
                # 使用 LLM 优化和过滤检索结果
                return await self._refine_with_llm(query, results)

        # Fallback: TF-IDF
        if len(self.chunks) > 0:
            try:
                from sklearn.feature_extraction.text import TfidfVectorizer
                from sklearn.metrics.pairwise import cosine_similarity
                import asyncio
                vec = await asyncio.to_thread(
                    lambda: TfidfVectorizer(analyzer="char", ngram_range=(1, 3),
                                            max_features=50000, sublinear_tf=True
                                           ).fit_transform(self.chunks + [query])
                )
                matrix = vec[:-1]
                qvec = vec[-1]
                sims = cosine_similarity(qvec, matrix).flatten()
                idxs = np.argsort(sims)[::-1][:top_k]
                results = [self.chunks[i] for i in idxs if sims[i] > 0.01]
                if results:
                    # 使用 LLM 优化和过滤检索结果
                    return await self._refine_with_llm(query, results)
            except Exception:
                pass

        return self._mock_retrieve(query, top_k)

    async def _refine_with_llm(self, query: str, chunks: list[str]) -> list[str]:
        """使用 LLM 对检索结果进行优化和相关性过滤"""
        from backend.app.core.llm import llm_service
        
        if not llm_service.llm or len(chunks) == 0:
            return chunks
        
        try:
            # 构建 prompt 让 LLM 评估相关性
            context = "\n\n".join([f"[片段{i+1}] {chunk}" for i, chunk in enumerate(chunks)])
            
            refine_prompt = f"""请分析以下检索片段与用户问题的相关性，只保留高度相关的片段。

用户问题：{query}

检索片段：
{context}

请输出相关片段的编号（如：1,3,5），如果不相关则输出：无"""
            
            # 调用 LLM 进行评估
            response, _ = await llm_service.chat(
                prompt=refine_prompt,
                context=None,
                system_prompt="你是一个专业的信息筛选助手，请准确判断文本相关性。"
            )
            
            # 解析 LLM 返回的相关片段编号
            import re
            numbers = re.findall(r'\d+', response)
            if numbers and response != "无":
                relevant_indices = [int(n) - 1 for n in numbers if 0 < int(n) <= len(chunks)]
                if relevant_indices:
                    refined_chunks = [chunks[i] for i in relevant_indices]
                    print(f"✅ LLM 优化检索结果: {len(chunks)} -> {len(refined_chunks)} 个片段")
                    return refined_chunks
            
            # 如果 LLM 无法解析，返回所有片段
            return chunks
            
        except Exception as e:
            print(f"⚠️ LLM 优化检索失败，使用原始结果: {e}")
            return chunks

    def _mock_retrieve(self, query: str, top_k: int) -> list[str]:
        knowledge = {
            "古塔": "古塔建于唐贞观年间（公元627-649年），七层八角，高约45米。塔身由青砖砌筑，每层都有精美的佛教石刻浮雕，塔内保存有唐代壁画残片。",
            "门票": "景区门票价格：成人票80元，学生票40元，60岁以上老人凭身份证免票。开放时间：每天8:00-18:00，17:00停止入园。",
            "历史": "景区始建于唐代，距今已有1300多年历史。相传一位高僧云游至此，见山水奇秀，便在此建寺修行，后逐渐发展成为著名风景区。",
            "美食": "景区美食街汇集了当地特色小吃：手工豆腐花、糯米糕、山野菌菇汤、特色烤鱼等。价格实惠，人均消费约30-50元。",
            "路线": "推荐文化探访路线：南门入口→古寺遗址→碑林→观景台→古塔，全程约2.5公里，步行约2小时。",
            "停车场": "景区南门和东门各有一个停车场，共提供500个车位，停车免费。旺季建议早到。",
            "宠物": "景区允许携带宠物入园，但需系牵引绳，大型犬需戴嘴套。",
        }
        results = []
        for key, content in knowledge.items():
            if key in query:
                results.append(content)
        if not results:
            results.append("景区总面积约3.5平方公里，分为历史文化区、自然生态区和休闲娱乐区三大板块。")
        return results[:top_k]

    async def rebuild_index(self):
        import shutil
        self.chunks.clear()
        self._embeddings = None
        if self._save_dir.exists():
            shutil.rmtree(self._save_dir)
        self._save_dir.mkdir(parents=True, exist_ok=True)


rag_service = RAGService()
