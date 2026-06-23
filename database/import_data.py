"""
数据库初始化与数据导入脚本

功能：
1. 创建 aihumanproject 数据库
2. 创建 tourist_behavior_data 表（游客行为数据）
3. 创建 attraction_details 表（景点详情）
4. 从 Excel/Word 文件导入数据
"""
import asyncio
import sys
from pathlib import Path
import pandas as pd
from docx import Document

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from sqlalchemy import create_engine, text
from backend.config import settings


def create_database_and_tables():
    """创建数据库和表结构"""
    print("=" * 80)
    print("🚀 开始创建数据库和表结构")
    print("=" * 80)
    
    # 读取SQL文件
    sql_file = project_root / "database" / "init_database.sql"
    if not sql_file.exists():
        print(f"❌ SQL文件不存在: {sql_file}")
        return False
    
    sql_content = sql_file.read_text(encoding='utf-8')
    
    try:
        # 连接MySQL（不指定数据库，先创建数据库）
        from sqlalchemy import create_engine
        import re
        
        # 解析数据库URL
        db_url_match = re.search(r'mysql\+aiomysql://([^:]+):([^@]+)@([^/]+)/(.+)', settings.DATABASE_URL)
        if db_url_match:
            username = db_url_match.group(1)
            password = db_url_match.group(2)
            host_port = db_url_match.group(3)
            dbname = db_url_match.group(4)
            
            # 创建不带数据库名的连接（用于创建数据库）
            base_url = f"mysql+pymysql://{username}:{password}@{host_port}"
            engine = create_engine(base_url)
            
            with engine.connect() as conn:
                from sqlalchemy import text
                # 执行每条SQL语句
                for statement in sql_content.split(';'):
                    statement = statement.strip()
                    if statement and not statement.startswith('--'):
                        try:
                            conn.execute(text(statement))
                        except Exception as e:
                            # 忽略某些错误（如表已存在）
                            if 'already exists' not in str(e).lower():
                                print(f"⚠️  警告: {e}")
                
                conn.commit()
            
            print("✅ 数据库和表结构创建成功！")
            return True
        else:
            print("❌ 无法解析数据库URL")
            return False
        
    except Exception as e:
        print(f"❌ 创建失败: {e}")
        import traceback
        traceback.print_exc()
        return False


async def import_tourist_behavior_data():
    """导入游客行为数据"""
    print("\n" + "=" * 80)
    print("📊 开始导入游客行为数据")
    print("=" * 80)
    
    # 原路径：excel_file = project_root / "backend" / "data" / "knowledge" / "景点景区旅游数据行为分析数据.xlsx"
    excel_file = project_root / "data" / "示范景区公开资料包" / "景点景区旅游数据行为分析数据.xlsx"
    
    if not excel_file.exists():
        print(f"❌ Excel文件不存在: {excel_file}")
        return False
    
    print(f"📁 读取文件: {excel_file.name}")
    print(f"📏 文件大小: {excel_file.stat().st_size / (1024*1024):.2f} MB")
    
    try:
        # 读取Excel
        df = pd.read_excel(excel_file)
        print(f"✅ 读取成功: {len(df)} 行 × {len(df.columns)} 列")
        
        # 显示列名
        print(f"\n📋 列名: {', '.join(df.columns.tolist())}")
        
        # 连接数据库（使用pymysql同步驱动）
        sync_db_url = settings.DATABASE_URL.replace('mysql+aiomysql://', 'mysql+pymysql://')
        engine = create_engine(sync_db_url)
        
        # 导入数据（分批插入，避免内存溢出）
        batch_size = 1000
        total_rows = len(df)
        inserted = 0
        
        print(f"\n🔄 开始导入数据（每批 {batch_size} 条）...")
        
        for i in range(0, total_rows, batch_size):
            batch_df = df.iloc[i:i+batch_size]
            
            # 重命名列以匹配数据库字段
            batch_df = batch_df.rename(columns={
                'tourist_id': 'tourist_id',
                'user_nickname': 'user_nickname',
                'age': 'age',
                'gender': 'gender',
                'attraction_name': 'attraction_name',
                'attraction_content': 'attraction_content',
                'attraction_type': 'attraction_type',
                'visit_date': 'visit_date',
                'stay_duration': 'stay_duration',
                'ticket_cost': 'ticket_cost',
                'food_cost': 'food_cost',
                'shopping_cost': 'shopping_cost',
                'transport_cost': 'transport_cost',
                'entertainment_cost': 'entertainment_cost',
                'total_cost': 'total_cost',
                'group_size': 'group_size',
                'satisfaction': 'satisfaction',
            })
            
            # 写入数据库
            batch_df.to_sql(
                'tourist_behavior_data',
                engine,
                if_exists='append',
                index=False,
                method='multi',
                chunksize=100
            )
            
            inserted += len(batch_df)
            progress = (inserted / total_rows) * 100
            print(f"  进度: {inserted}/{total_rows} ({progress:.1f}%)")
        
        print(f"\n✅ 游客行为数据导入完成！共 {inserted} 条记录")
        return True
        
    except Exception as e:
        print(f"❌ 导入失败: {e}")
        import traceback
        traceback.print_exc()
        return False


async def import_attraction_details():
    """导入景点详情数据"""
    print("\n" + "=" * 80)
    print("📋 开始导入景点详情数据")
    print("=" * 80)
    
    # 原逻辑（已注释，保留参考）：
    # docx_file = project_root / "backend" / "data" / "knowledge" / "灵山胜境 景点结构化数据集.docx"
    # if not docx_file.exists():
    #     print(f"❌ Word文件不存在: {docx_file}")
    #     return False
    # print(f"📁 读取文件: {docx_file.name}")
    # try:
    #     doc = Document(docx_file)
    #     print(f"✅ 读取成功: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")
    #     attractions = []
    #     for table_idx, table in enumerate(doc.tables):
    #         print(f"\n📊 处理表格 {table_idx + 1}: {len(table.rows)} 行 × {len(table.columns)} 列")
    #         if table.rows:
    #             headers = [cell.text.strip() for cell in table.rows[0].cells]
    #             print(f"   表头: {headers[:5]}...")
    #         for row_idx, row in enumerate(table.rows[1:], 1):
    #             cells = [cell.text.strip() for cell in row.cells]
    #             if len(cells) >= 11:
    #                 attraction = {
    #                     'scenic_area': cells[0],
    #                     'attraction_id': cells[1],
    #                     'attraction_name': cells[2],
    #                     'location': cells[3],
    #                     'building_params': cells[4],
    #                     'core_function': cells[5],
    #                     'cultural_meaning': cells[6],
    #                     'detailed_description': cells[7],
    #                     'highlights': cells[8],
    #                     'performance_info': cells[9],
    #                     'remarks': cells[10],
    #                     'data_source': docx_file.name,
    #                 }
    #                 attractions.append(attraction)
    #     print(f"\n✅ 提取到 {len(attractions)} 个景点信息")
    #     df = pd.DataFrame(attractions)

    # 新逻辑：直接读取清理好的 CSV 数据文件
    csv_file = project_root / "data" / "示范景区公开资料包" / "景点详情.csv"
    
    if not csv_file.exists():
        print(f"❌ CSV文件不存在: {csv_file}")
        return False
        
    print(f"📁 读取文件: {csv_file.name}")
    
    try:
        # 使用 utf-8-sig 编码自动处理中文 BOM 头
        df = pd.read_csv(csv_file, encoding='utf-8-sig')
        print(f"✅ 读取成功: {len(df)} 行 × {len(df.columns)} 列")
        
        # 重命名 CSV 列以完全匹配数据库表的属性字段
        df = df.rename(columns={
            '景区名称': 'scenic_area',
            '景点ID': 'attraction_id',
            '景点名称': 'attraction_name',
            '具体位置': 'location',
            '建筑/景观参数': 'building_params',
            '核心功能': 'core_function',
            '文化内涵': 'cultural_meaning',
            '详细介绍': 'detailed_description',
            '游玩亮点': 'highlights',
            '演艺/开放信息': 'performance_info',
            '备注': 'remarks'
        })
        df['data_source'] = csv_file.name
        
        # 连接数据库（使用pymysql同步驱动）
        sync_db_url = settings.DATABASE_URL.replace('mysql+aiomysql://', 'mysql+pymysql://')
        engine = create_engine(sync_db_url)
        
        # 写入数据库
        df.to_sql(
            'attraction_details',
            engine,
            if_exists='append',
            index=False,
            method='multi',
            chunksize=50
        )
        
        print(f"✅ 景点详情数据导入完成！共 {len(df)} 条记录")
        return True
        
    except Exception as e:
        print(f"❌ 导入失败: {e}")
        import traceback
        traceback.print_exc()
        return False


async def verify_import():
    """验证数据导入"""
    print("\n" + "=" * 80)
    print("🔍 验证数据导入")
    print("=" * 80)
    
    try:
        # 连接数据库（使用pymysql同步驱动）
        sync_db_url = settings.DATABASE_URL.replace('mysql+aiomysql://', 'mysql+pymysql://')
        engine = create_engine(sync_db_url)
        
        with engine.connect() as conn:
            # 检查tourist_behavior_data表
            result = conn.execute(text("SELECT COUNT(*) as count FROM tourist_behavior_data"))
            count = result.fetchone()[0]
            print(f"📊 tourist_behavior_data: {count} 条记录")
            
            # 检查attraction_details表
            result = conn.execute(text("SELECT COUNT(*) as count FROM attraction_details"))
            count = result.fetchone()[0]
            print(f"📋 attraction_details: {count} 条记录")
            
            # 显示示例数据
            print(f"\n📝 游客行为数据示例:")
            result = conn.execute(text("SELECT tourist_id, user_nickname, attraction_name, visit_date, total_cost FROM tourist_behavior_data LIMIT 3"))
            for row in result:
                print(f"   {row}")
            
            print(f"\n📝 景点详情示例:")
            result = conn.execute(text("SELECT attraction_id, attraction_name, scenic_area FROM attraction_details LIMIT 3"))
            for row in result:
                print(f"   {row}")
        
        print("\n✅ 数据验证完成！")
        return True
        
    except Exception as e:
        print(f"❌ 验证失败: {e}")
        return False


async def main():
    """主函数"""
    print("\n" + "=" * 80)
    print("🎯 数据库初始化与数据导入工具")
    print("=" * 80)
    
    # 步骤1: 创建数据库和表
    if not create_database_and_tables():
        print("\n❌ 数据库创建失败，终止操作")
        return
    
    # 步骤2: 导入游客行为数据
    if not await import_tourist_behavior_data():
        print("\n⚠️  游客行为数据导入失败，继续下一步")
    
    # 步骤3: 导入景点详情数据
    if not await import_attraction_details():
        print("\n⚠️  景点详情数据导入失败")
    
    # 步骤4: 验证导入
    await verify_import()
    
    print("\n" + "=" * 80)
    print("✨ 所有操作完成！")
    print("=" * 80)


if __name__ == "__main__":
    asyncio.run(main())
